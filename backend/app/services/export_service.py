# Servicio de exportación: renderiza el apunte a TXT, DOCX o PDF.
import io
import re
from xml.sax.saxutils import escape

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
    ListItem,
    Preformatted,
    Table,
    TableStyle,
)
from docx import Document
from docx.shared import Pt, RGBColor, Inches

FORGE_BLUE = "#2563eb"
NEUTRAL = "#404040"

# Tipo MIME por formato de exportación.
MIMETYPES = {
    "txt": "text/plain; charset=utf-8",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _bullet_text(item) -> str:
    if isinstance(item, str):
        return item
    if isinstance(item, dict):
        return item.get("text") or item.get("latex") or ""
    return str(item)


# Sanea el título para usarlo como nombre de fichero.
def safe_filename(title: str, ext: str) -> str:
    base = (title or "apunte").strip()
    base = re.sub(r"[^\w\s\-.áéíóúàèìòùçñÁÉÍÓÚÀÈÌÒÙÇÑ]", "", base, flags=re.UNICODE)
    base = re.sub(r"\s+", "_", base).strip("_")
    return f"{base or 'apunte'}.{ext}"


# Despachador: elige el generador según el formato pedido.
def render(apunte: dict, fmt: str):
    if fmt == "txt":
        return to_txt(apunte), MIMETYPES["txt"], "txt"
    if fmt == "docx":
        return to_docx(apunte), MIMETYPES["docx"], "docx"
    return to_pdf(apunte), MIMETYPES["pdf"], "pdf"


# Versión en texto plano del apunte.
def to_txt(apunte: dict) -> bytes:
    lines = []
    title = apunte.get("title") or "Apunte"
    lines.append(title)
    lines.append("=" * len(title))
    lines.append("")

    tags = apunte.get("tags") or []
    if tags:
        lines.append("Etiquetas: " + ", ".join(str(t) for t in tags))
        lines.append("")

    if apunte.get("summary"):
        lines.append("RESUMEN")
        lines.append(apunte["summary"])
        lines.append("")

    structure = apunte.get("structure") or {}
    for i, sec in enumerate(structure.get("sections", []), start=1):
        lines.append(f"{i}. {sec.get('heading', '')}")
        lines.append("")
        for block in sec.get("blocks", []):
            # Cada tipo de bloque (párrafo, lista, fórmula, cita) se vuelca distinto.
            btype = block.get("type")
            if btype == "paragraph":
                lines.append(block.get("text", ""))
                lines.append("")
            elif btype == "bullet_list":
                for it in block.get("items", []):
                    lines.append(f"  - {_bullet_text(it)}")
                lines.append("")
            elif btype == "formula":
                lines.append(f"    {block.get('latex', '')}")
                lines.append("")
            elif btype == "quote":
                src = f" ({block.get('source')})" if block.get("source") else ""
                lines.append(f'  "{block.get("text", "")}"{src}')
                lines.append("")

    lines.append("")
    lines.append("Generado con NoteForge")
    return ("\n".join(lines)).encode("utf-8")


# Versión Word con estilos básicos (python-docx).
def to_docx(apunte: dict) -> bytes:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_heading(apunte.get("title") or "Apunte", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    tags = apunte.get("tags") or []
    if tags:
        p = doc.add_paragraph()
        run = p.add_run(" · ".join(str(t) for t in tags))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x71, 0x71, 0x71)

    if apunte.get("summary"):
        h = doc.add_heading("Resumen", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        sp = doc.add_paragraph(apunte["summary"])
        sp.paragraph_format.left_indent = Inches(0.15)

    structure = apunte.get("structure") or {}
    for i, sec in enumerate(structure.get("sections", []), start=1):
        h = doc.add_heading(f"{i}. {sec.get('heading', '')}", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        for block in sec.get("blocks", []):
            btype = block.get("type")
            if btype == "paragraph":
                doc.add_paragraph(block.get("text", ""))
            elif btype == "bullet_list":
                for it in block.get("items", []):
                    doc.add_paragraph(_bullet_text(it), style="List Bullet")
            elif btype == "formula":
                p = doc.add_paragraph()
                run = p.add_run(block.get("latex", ""))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            elif btype == "quote":
                p = doc.add_paragraph(style="Quote")
                src = f"  — {block.get('source')}" if block.get("source") else ""
                p.add_run(f'"{block.get("text", "")}"{src}')

    footer = doc.add_paragraph()
    fr = footer.add_run("Generado con NoteForge")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xA3, 0xA3, 0xA3)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Versión PDF maquetada con ReportLab.
def to_pdf(apunte: dict) -> bytes:
    buf = io.BytesIO()
    docpdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=22 * mm,
        rightMargin=22 * mm,
        topMargin=22 * mm,
        bottomMargin=20 * mm,
        title=apunte.get("title") or "Apunte",
    )

    base = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "NFTitle", parent=base["Title"], fontName="Helvetica-Bold",
        fontSize=24, leading=28, textColor=HexColor("#111827"), spaceAfter=4,
    )
    tag_style = ParagraphStyle(
        "NFTags", parent=base["Normal"], fontName="Helvetica",
        fontSize=9, textColor=HexColor("#737373"), spaceAfter=14,
    )
    h_style = ParagraphStyle(
        "NFHeading", parent=base["Heading2"], fontName="Helvetica-Bold",
        fontSize=15, leading=19, textColor=HexColor("#111827"),
        spaceBefore=16, spaceAfter=8,
    )
    summary_style = ParagraphStyle(
        "NFSummary", parent=base["Normal"], fontName="Helvetica",
        fontSize=10.5, leading=15, textColor=HexColor(NEUTRAL), spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "NFBody", parent=base["Normal"], fontName="Helvetica",
        fontSize=11, leading=16, alignment=TA_LEFT, spaceAfter=8,
    )
    quote_style = ParagraphStyle(
        "NFQuote", parent=base["Normal"], fontName="Helvetica-Oblique",
        fontSize=10.5, leading=15, leftIndent=10, textColor=HexColor("#525252"),
        spaceAfter=8,
    )
    formula_style = ParagraphStyle(
        "NFFormula", parent=base["Code"], fontName="Courier",
        fontSize=9.5, leading=13, backColor=HexColor("#f5f5f5"),
        borderPadding=6, spaceAfter=8,
    )

    story = []
    story.append(Paragraph(escape(apunte.get("title") or "Apunte"), title_style))

    tags = apunte.get("tags") or []
    if tags:
        story.append(Paragraph(escape(" · ".join(str(t) for t in tags)), tag_style))

    if apunte.get("summary"):
        box = Table(
            [[Paragraph(escape(apunte["summary"]), summary_style)]],
            colWidths=[docpdf.width],
        )
        box.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), HexColor("#eff6ff")),
            ("BOX", (0, 0), (-1, -1), 0.5, HexColor("#bfdbfe")),
            ("LEFTPADDING", (0, 0), (-1, -1), 12),
            ("RIGHTPADDING", (0, 0), (-1, -1), 12),
            ("TOPPADDING", (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ]))
        story.append(box)
        story.append(Spacer(1, 16))

    structure = apunte.get("structure") or {}
    for i, sec in enumerate(structure.get("sections", []), start=1):
        story.append(Paragraph(f"{i}. {escape(sec.get('heading', ''))}", h_style))
        for block in sec.get("blocks", []):
            btype = block.get("type")
            if btype == "paragraph":
                story.append(Paragraph(escape(block.get("text", "")), body_style))
            elif btype == "bullet_list":
                items = [
                    ListItem(Paragraph(escape(_bullet_text(it)), body_style), leftIndent=14)
                    for it in block.get("items", [])
                ]
                if items:
                    story.append(ListFlowable(items, bulletType="bullet", start="•", spaceAfter=8))
            elif btype == "formula":
                story.append(Preformatted(block.get("latex", ""), formula_style))
            elif btype == "quote":
                src = f"  — {escape(str(block.get('source')))}" if block.get("source") else ""
                story.append(Paragraph(f'"{escape(block.get("text", ""))}"{src}', quote_style))

    docpdf.build(story)
    return buf.getvalue()